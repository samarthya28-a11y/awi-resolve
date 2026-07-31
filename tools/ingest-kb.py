#!/usr/bin/env python
"""
AWI Resolve - knowledge-base ingestion.

Turns product documentation (PDF / DOCX / TXT) into a searchable knowledge base
the AI technician can look things up in, instead of dumping whole 200-page
manuals into its context.

Each document is split into overlapping chunks with their page numbers, so the
AI gets a short, citable excerpt ("Server manual v9, p.142") rather than a book.

Usage:
    python tools/ingest-kb.py "<source folder>" [--tag gespage] [--max-mb 40]

Output: playbooks/kb/<slug>.json   (one file per source document)
"""
import argparse
import json
import pathlib
import re
import sys
import unicodedata

CHUNK_CHARS = 1800      # ~450 tokens: big enough to hold a procedure
CHUNK_OVERLAP = 250     # keeps a step from being cut in half
MIN_CHUNK_CHARS = 120   # skip near-empty pages (covers, dividers)


def slugify(text: str) -> str:
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode()
    text = re.sub(r"[^\w\s-]", "", text).strip().lower()
    return re.sub(r"[\s_-]+", "-", text)[:80]


def clean(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t ]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)     # collapse blank runs
    return text.strip()


def read_pdf(path: pathlib.Path):
    """Yield (page_number, text). Uses PyMuPDF."""
    import fitz
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, start=1):
            yield i, page.get_text("text") or ""


def read_docx(path: pathlib.Path):
    try:
        from docx import Document
    except ImportError:
        return
    doc = Document(path)
    body = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            body += "\n" + " | ".join(c.text.strip() for c in row.cells)
    yield 1, body


def read_txt(path: pathlib.Path):
    yield 1, path.read_text(encoding="utf-8", errors="replace")


READERS = {".pdf": read_pdf, ".docx": read_docx,
           ".txt": read_txt, ".md": read_txt, ".log": read_txt}


def chunk_pages(pages):
    """Chunk text while remembering which page each chunk came from."""
    chunks, buf, buf_start_page = [], "", None
    for page_no, raw in pages:
        text = clean(raw)
        if len(text) < 20:
            continue
        if buf_start_page is None:
            buf_start_page = page_no
        buf += ("\n\n" if buf else "") + text
        while len(buf) >= CHUNK_CHARS:
            piece, buf = buf[:CHUNK_CHARS], buf[CHUNK_CHARS - CHUNK_OVERLAP:]
            if len(piece.strip()) >= MIN_CHUNK_CHARS:
                chunks.append({"page": buf_start_page, "text": piece.strip()})
            buf_start_page = page_no
    if len(buf.strip()) >= MIN_CHUNK_CHARS:
        chunks.append({"page": buf_start_page or 1, "text": buf.strip()})
    return chunks


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("source", help="folder containing the documentation")
    ap.add_argument("--tag", default="", help="product tag, e.g. gespage")
    ap.add_argument("--max-mb", type=float, default=40.0, help="skip files larger than this")
    ap.add_argument("--include", default="", help="only files whose path matches this regex (case-insensitive)")
    args = ap.parse_args()

    src = pathlib.Path(args.source)
    if not src.is_dir():
        print(f"ERROR: not a folder: {src}", file=sys.stderr)
        return 1

    out_dir = pathlib.Path(__file__).resolve().parent.parent / "playbooks" / "kb"
    out_dir.mkdir(parents=True, exist_ok=True)

    include = re.compile(args.include, re.I) if args.include else None
    total_chunks = written = skipped = 0

    for path in sorted(src.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in READERS:
            continue
        rel = str(path.relative_to(src))
        if include and not include.search(rel):
            continue
        if path.stat().st_size > args.max_mb * 1024 * 1024:
            print(f"  skip (too big): {rel}")
            skipped += 1
            continue
        try:
            chunks = chunk_pages(READERS[path.suffix.lower()](path))
        except Exception as exc:                      # noqa: BLE001 - report and continue
            print(f"  skip (unreadable): {rel} -- {exc}")
            skipped += 1
            continue
        if not chunks:
            print(f"  skip (no text - probably scanned): {rel}")
            skipped += 1
            continue

        slug = slugify(f"{args.tag}-{path.stem}" if args.tag else path.stem)
        doc = {
            "id": slug,
            "title": path.stem,
            "source": rel,
            "tag": args.tag,
            "chunkCount": len(chunks),
            "chunks": chunks,
        }
        (out_dir / f"{slug}.json").write_text(json.dumps(doc, ensure_ascii=False), encoding="utf-8")
        total_chunks += len(chunks)
        written += 1
        print(f"  ok: {rel}  ({len(chunks)} chunks)")

    print(f"\nIngested {written} document(s), {total_chunks} chunks -> {out_dir}")
    if skipped:
        print(f"Skipped {skipped} file(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
