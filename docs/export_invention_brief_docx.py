from docx import Document
from docx.shared import Pt, Inches, RGBColor
from pathlib import Path

out = Path(__file__).with_name("AWI-Resolve-Invention-Brief.docx")
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)


def add_para(text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x15, 0x21, 0x3B)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(3)
    return p


title = doc.add_heading("Invention Brief — AWI Resolve", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x15, 0x21, 0x3B)
    run.font.size = Pt(16)

add_para(
    "Autonomous, consent-gated IT support agent with dual-gated privileged "
    "execution and pinned software deployment",
    italic=True,
    size=11,
)

meta = [
    ("Prepared for:", "Patent counsel (intake only — not a patent application)"),
    ("Applicant (indicative):", "Alpha Web Innovations Private Limited"),
    ("Product:", "AWI Resolve (on-device Windows agent + orchestrator)"),
    ("Date:", "17 August 2026"),
    ("Status:", "Confidential — for attorney review"),
]
for k, v in meta:
    p = doc.add_paragraph()
    r1 = p.add_run(k + " ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(v)
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

add_heading("1. Problem", 1)
add_para(
    "Organisations want AI-driven IT support on end-user PCs, but existing "
    "approaches fail one or more safety requirements:"
)
add_bullet(
    "Unsupervised shell / RMM can change the machine without the end user "
    "seeing the exact action."
)
add_bullet(
    "Chatbots guide users but cannot safely apply fixes or install software "
    "with auditability."
)
add_bullet(
    "“Attach an installer / URL” workflows let end users smuggle unvetted "
    "binaries into automated install paths."
)
add_bullet(
    "Privileged AI modes (free-form PowerShell) are either always on (too risky) "
    "or absent (too weak), with no organisational second gate."
)
add_bullet(
    "Customer-supplied manuals either never reach the model mid-session, or "
    "are treated as instructions that could jailbreak the agent."
)
add_para(
    "Needed: a system that does diagnose and fix, can deploy software, and can "
    "run broader commands when authorised — without letting end-user attachments "
    "or a single licence flag become a security bypass."
)

add_heading("2. Technical solution (summary)", 1)
add_para(
    "AWI Resolve runs a local Windows agent (diagnostics, consent UI, "
    "install/execution) and a support orchestrator (AI planner, catalogues, "
    "licensing, session state)."
)
add_para("Core mechanism — layered authority for change:", bold=True)

table = doc.add_table(rows=6, cols=2)
table.style = "Table Grid"
headers = ("Layer", "What it controls")
rows = [
    (
        "A. Tool allowlist",
        "Default plans only invoke vetted tools (diagnostics + named fixes). "
        "No free-form shell.",
    ),
    (
        "B. End-user consent",
        "Any non-trivial change (service restart, deploy, PowerShell) shows the "
        "exact action/command; decline/timeout escalates.",
    ),
    (
        "C. Global + org software catalogues",
        "Auto-install only via productId → pinned HTTPS URL + SHA-256 "
        "(and optional settings). End-user chat attachments cannot supply "
        "install URLs or unlock catalogue installs.",
    ),
    (
        "D. Dual-gated Full IT Support",
        "run_powershell (and off-catalogue official installs via it) require "
        "both a Full licence capability and an org-admin “allow Full IT Support” "
        "toggle. Exact command shown for Yes/No; runtime/output capped; "
        "illegitimate asks refused in policy.",
    ),
    (
        "E. Untrusted document channel",
        "Customer PDFs/screenshots are queued per device, injected into the "
        "same conversation (including follow-ups) as marked untrusted reference "
        "data — usable for guidance, not as install authority.",
    ),
]
table.rows[0].cells[0].text = headers[0]
table.rows[0].cells[1].text = headers[1]
for i, (a, b) in enumerate(rows, start=1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"
for cell in table.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

doc.add_paragraph()
add_para(
    "Session outputs include a structured report (what was checked, changed, "
    "declined, not done) for handoff."
)

add_heading(
    "3. Claims-style bullets (for counsel — illustrative, not formal claims)",
    1,
)
claims = [
    "A method of autonomous endpoint IT support comprising: performing "
    "read-only diagnostics via an on-device agent; proposing machine-changing "
    "actions only through a fixed tool interface; and gating each such action "
    "on explicit end-user consent that displays the precise action or command text.",
    "The method of (1), wherein software installation is permitted only by "
    "resolving a catalogue or organisation-library product identifier to a "
    "server-side pinned download location and cryptographic checksum, and wherein "
    "content attached by the end user in chat is excluded from supplying install "
    "locations or product identifiers for automatic installation.",
    "The method of (1)–(2), further comprising enabling a privileged execution "
    "tool (e.g. arbitrary PowerShell) only when both (i) a licence/capability "
    "flag for full support and (ii) an organisation-administrator policy flag "
    "are asserted, and presenting the exact command string for consent before "
    "execution under resource caps.",
    "The method of (3), wherein after catalogue and organisation-library checks "
    "fail to match a user-requested product, the privileged tool may install from "
    "a vendor official HTTPS source or package manager, subject to the dual gate "
    "and consent of (3).",
    "A system comprising an on-device agent and an orchestrator that maintain "
    "per-device pending attachments; extract text from uploaded documents; and "
    "inject those attachments into an ongoing multi-turn support session as "
    "delimited untrusted reference data without treating them as executable "
    "instructions or install authorisations.",
    "The system of (1)–(5), further generating a session record enumerating "
    "diagnostics performed, actions applied, actions declined or timed out, and "
    "items explicitly not done, for human escalation.",
    "A computer-readable medium storing instructions that, when executed, cause "
    "a computing system to perform any of (1)–(6).",
]
for c in claims:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(c)
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)

add_heading("4. Embodiments / implementation notes (non-limiting)", 1)
add_bullet(
    "Windows 10/11 agent UI with paperclip, clipboard paste, and drag-and-drop "
    "for images/docs."
)
add_bullet(
    "Org software library: admin-stored manual text + HTTPS link + SHA-256; "
    "deploy tool accepts only productId."
)
add_bullet(
    "Licensing plans (e.g. trial / standard / pro / full / time-bounded pass) "
    "map to capability sets including fullSupport."
)
add_bullet(
    "Orchestrator refuses off-allowlist tools at the agent; safety probe can "
    "verify refusal."
)

add_heading("5. Prior art to search (suggested)", 1)
add_para(
    "AI IT helpdesks; RMM with script approval; Intune/winget automation; "
    "“human-in-the-loop” LLM tool use; software catalogues with hash verification; "
    "DLP/untrusted document wrapping for LLMs."
)

add_heading("6. Inventorship / filing notes for counsel", 1)
add_bullet(
    "Confirm inventors (engineering contributors to layers C–E and dual gate)."
)
add_bullet(
    "Check public disclosures (website, trials, GitHub) vs novelty deadlines "
    "(India / PCT)."
)
add_bullet(
    "Consider provisional + later PCT; parallel trademark for “AWI Resolve”."
)
add_bullet(
    "Preserve unpublished implementation details as trade secret where not claimed."
)

add_para(
    "End of invention brief. Expand into formal specification only under "
    "attorney direction.",
    italic=True,
)

doc.save(out)
print(f"Wrote {out}")
print(f"Size bytes: {out.stat().st_size}")
